"""
Document processing pipeline for handling files of any size.

Supports PDF, DOCX, XLSX with intelligent chunking and
optional embedding for large documents.
"""

import asyncio
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field


@dataclass
class Chunk:
    """A chunk of document content."""
    text: str
    page: Optional[int] = None
    section: Optional[str] = None
    index: int = 0


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    path: str
    filename: str
    file_type: str
    total_pages: int = 0
    total_chars: int = 0
    chunks: List[Chunk] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    error: Optional[str] = None

    @property
    def is_small(self) -> bool:
        """Documents under 15KB can be injected directly into context."""
        return self.total_chars < 15000

    @property
    def full_text(self) -> str:
        """Get full document text."""
        return "\n\n".join(c.text for c in self.chunks)


# Size threshold: below this, inject directly into context
DIRECT_CONTEXT_THRESHOLD = 15000


class DocumentProcessor:
    """
    Processes documents of any size intelligently.

    For small docs (< 15KB): Direct context injection (fast)
    For large docs: Chunk → embed → retrieve relevant sections per query
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    async def process(self, file_path: str) -> ProcessedDocument:
        """
        Process a document file.

        Supports: PDF, DOCX, XLSX, TXT, MD, CSV, JSON
        """
        path = Path(file_path)
        if not path.exists():
            return ProcessedDocument(
                path=file_path,
                filename=path.name,
                file_type="unknown",
                error=f"File not found: {file_path}",
            )

        ext = path.suffix.lower()
        processors = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".xlsx": self._process_xlsx,
            ".xls": self._process_xlsx,
            ".csv": self._process_csv,
            ".txt": self._process_text,
            ".md": self._process_text,
            ".json": self._process_json,
            ".py": self._process_text,
            ".js": self._process_text,
            ".ts": self._process_text,
            ".yaml": self._process_text,
            ".yml": self._process_text,
        }

        processor = processors.get(ext, self._process_text)

        try:
            return await asyncio.to_thread(processor, path)
        except Exception as e:
            return ProcessedDocument(
                path=file_path,
                filename=path.name,
                file_type=ext.lstrip("."),
                error=str(e),
            )

    def _process_pdf(self, path: Path) -> ProcessedDocument:
        """Process PDF file page by page."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            # Fallback to basic text extraction
            return self._process_text(path)

        doc = fitz.open(str(path))
        chunks = []
        total_chars = 0

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()

            if text.strip():
                # Split large pages into smaller chunks
                page_chunks = self.chunk_text(text, page=page_num + 1)
                chunks.extend(page_chunks)
                total_chars += len(text)

        doc.close()

        return ProcessedDocument(
            path=str(path),
            filename=path.name,
            file_type="pdf",
            total_pages=len(doc) if hasattr(doc, '__len__') else 0,
            total_chars=total_chars,
            chunks=chunks,
            metadata={"processor": "pymupdf"},
        )

    def _process_docx(self, path: Path) -> ProcessedDocument:
        """Process DOCX file with structure preservation."""
        try:
            from docx import Document
        except ImportError:
            return self._process_text(path)

        doc = Document(str(path))
        sections = []
        current_section = ""
        current_heading = ""

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                # Save previous section
                if current_section.strip():
                    sections.append((current_heading, current_section.strip()))
                current_heading = para.text
                current_section = ""
            else:
                current_section += para.text + "\n"

        # Save last section
        if current_section.strip():
            sections.append((current_heading, current_section.strip()))

        # Process tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                table_text = "\n".join(rows)
                sections.append(("Table", table_text))

        # Create chunks from sections
        chunks = []
        total_chars = 0
        for heading, content in sections:
            section_chunks = self.chunk_text(content, section=heading or None)
            chunks.extend(section_chunks)
            total_chars += len(content)

        return ProcessedDocument(
            path=str(path),
            filename=path.name,
            file_type="docx",
            total_chars=total_chars,
            chunks=chunks,
            metadata={"sections": len(sections)},
        )

    def _process_xlsx(self, path: Path) -> ProcessedDocument:
        """Process Excel file with all sheets."""
        try:
            import pandas as pd
        except ImportError:
            return ProcessedDocument(
                path=str(path),
                filename=path.name,
                file_type="xlsx",
                error="pandas not installed",
            )

        chunks = []
        total_chars = 0

        xls = pd.ExcelFile(str(path))
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            # Convert to markdown table
            md_table = df.to_markdown(index=False)
            if md_table:
                sheet_chunks = self.chunk_text(md_table, section=f"Sheet: {sheet_name}")
                chunks.extend(sheet_chunks)
                total_chars += len(md_table)

        return ProcessedDocument(
            path=str(path),
            filename=path.name,
            file_type="xlsx",
            total_chars=total_chars,
            chunks=chunks,
            metadata={"sheets": len(xls.sheet_names)},
        )

    def _process_csv(self, path: Path) -> ProcessedDocument:
        """Process CSV file."""
        try:
            import pandas as pd
            df = pd.read_csv(str(path))
            text = df.to_markdown(index=False)
        except ImportError:
            text = path.read_text(errors='replace')

        chunks = self.chunk_text(text or "")
        return ProcessedDocument(
            path=str(path),
            filename=path.name,
            file_type="csv",
            total_chars=len(text or ""),
            chunks=chunks,
        )

    def _process_json(self, path: Path) -> ProcessedDocument:
        """Process JSON file."""
        import json

        content = path.read_text(errors='replace')
        try:
            data = json.loads(content)
            text = json.dumps(data, indent=2)
        except json.JSONDecodeError:
            text = content

        chunks = self.chunk_text(text)
        return ProcessedDocument(
            path=str(path),
            filename=path.name,
            file_type="json",
            total_chars=len(text),
            chunks=chunks,
        )

    def _process_text(self, path: Path) -> ProcessedDocument:
        """Process plain text file."""
        content = path.read_text(errors='replace')
        chunks = self.chunk_text(content)

        return ProcessedDocument(
            path=str(path),
            filename=path.name,
            file_type=path.suffix.lstrip(".") or "txt",
            total_chars=len(content),
            chunks=chunks,
        )

    def chunk_text(
        self,
        text: str,
        page: Optional[int] = None,
        section: Optional[str] = None,
    ) -> List[Chunk]:
        """
        Split text into overlapping chunks.

        Uses paragraph boundaries when possible to keep semantic coherence.
        """
        if not text or not text.strip():
            return []

        # If text fits in one chunk, return as-is
        if len(text) <= self.chunk_size:
            return [Chunk(text=text, page=page, section=section, index=0)]

        chunks = []
        # Try to split on paragraph boundaries
        paragraphs = text.split("\n\n")

        current = ""
        idx = 0

        for para in paragraphs:
            if len(current) + len(para) + 2 > self.chunk_size and current:
                chunks.append(Chunk(
                    text=current.strip(),
                    page=page,
                    section=section,
                    index=idx,
                ))
                idx += 1
                # Keep overlap from end of current chunk
                if self.chunk_overlap > 0:
                    current = current[-self.chunk_overlap:] + "\n\n" + para
                else:
                    current = para
            else:
                current = current + "\n\n" + para if current else para

        # Don't forget last chunk
        if current.strip():
            chunks.append(Chunk(
                text=current.strip(),
                page=page,
                section=section,
                index=idx,
            ))

        return chunks

    async def embed_and_store(
        self,
        doc: ProcessedDocument,
        collection_name: str = None,
    ) -> str:
        """
        Embed document chunks into ChromaDB for retrieval.

        Returns collection name for later querying.
        """
        col_name = collection_name or f"doc_{Path(doc.filename).stem}"

        try:
            import chromadb
            client = chromadb.Client()
            collection = client.get_or_create_collection(
                name=col_name,
                metadata={"hnsw:space": "cosine"},
            )

            documents = []
            ids = []
            metadatas = []

            for i, chunk in enumerate(doc.chunks):
                documents.append(chunk.text)
                ids.append(f"{col_name}_chunk_{i}")
                metadatas.append({
                    "page": chunk.page or 0,
                    "section": chunk.section or "",
                    "index": chunk.index,
                    "source": doc.filename,
                })

            if documents:
                # Add in batches
                batch_size = 100
                for start in range(0, len(documents), batch_size):
                    end = min(start + batch_size, len(documents))
                    await asyncio.to_thread(
                        collection.add,
                        documents=documents[start:end],
                        ids=ids[start:end],
                        metadatas=metadatas[start:end],
                    )

            return col_name

        except ImportError:
            return ""
        except Exception as e:
            import logging
            logging.getLogger("mike.docprocessor").warning(f"Embedding error: {e}")
            return ""

    async def query_document(
        self,
        collection_name: str,
        question: str,
        n_results: int = 5,
    ) -> str:
        """
        Query embedded document chunks.

        Returns relevant sections concatenated.
        """
        try:
            import chromadb
            client = chromadb.Client()
            collection = client.get_collection(collection_name)

            results = await asyncio.to_thread(
                collection.query,
                query_texts=[question],
                n_results=min(n_results, collection.count()),
            )

            if results and results.get("documents"):
                sections = results["documents"][0]
                return "\n\n---\n\n".join(sections)

        except Exception as e:
            import logging
            logging.getLogger("mike.docprocessor").warning(f"Query error: {e}")

        return ""

    async def process_for_context(
        self,
        file_path: str,
        query: str = "",
    ) -> str:
        """
        Process a document and return content suitable for LLM context.

        Small docs: returns full text
        Large docs: embeds and retrieves relevant sections

        Args:
            file_path: Path to the document
            query: Optional query for large document retrieval

        Returns:
            Document content string
        """
        doc = await self.process(file_path)

        if doc.error:
            return f"Error processing {doc.filename}: {doc.error}"

        if doc.is_small:
            # Direct injection for small documents
            return f"Document: {doc.filename}\n\n{doc.full_text}"

        # Large document: embed and retrieve
        col_name = await self.embed_and_store(doc)
        if col_name and query:
            relevant = await self.query_document(col_name, query)
            if relevant:
                return (
                    f"Document: {doc.filename} ({doc.total_chars:,} chars, {len(doc.chunks)} sections)\n"
                    f"Relevant sections for your query:\n\n{relevant}"
                )

        # Fallback: return first N chunks
        max_chars = DIRECT_CONTEXT_THRESHOLD
        text = ""
        for chunk in doc.chunks:
            if len(text) + len(chunk.text) > max_chars:
                break
            text += chunk.text + "\n\n"

        return (
            f"Document: {doc.filename} ({doc.total_chars:,} chars total, showing first {len(text):,} chars)\n\n"
            f"{text}\n\n... [Document truncated. Ask specific questions for targeted retrieval.]"
        )
